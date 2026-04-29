"""
API для аналитической системы демографии РФ
"""
import os
import re
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any

import numpy as np
from xml.sax.saxutils import escape
from fastapi import FastAPI, HTTPException
from fastapi.responses import PlainTextResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from docx import Document

# Библиотеки для генерации PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4

import re
from xml.sax.saxutils import escape
from reportlab.platypus import Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.fonts import addMapping
# Импорт внутренних модулей
from task1_monitoring import PopulationMonitor
from task2_forecasting import PopulationForecaster
from task3_ai_analytics import AIAnalytics
from database import (
    init_db, save_forecast, get_all_forecasts,
    get_forecast_by_id, get_forecasts_by_city, delete_forecast
)
from task4_llm_analytics import LLMAnalytics


BASE_DIR = Path(__file__).resolve().parent
REPORTS_DIR = BASE_DIR / "reports"
REPORTS_DIR.mkdir(exist_ok=True)
# Рекомендуется использовать DejaVuSans.ttf для лучшей поддержки кириллицы
FONT_PATH = BASE_DIR / "fonts" / "DejaVuSans.ttf"
# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

FONT_NAME = "Helvetica"

def register_fonts():
    global FONT_NAME
    # Путь к шрифту Arial в Windows
    font_path = "C:/Windows/Fonts/arial.ttf"
    try:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont("Arial", font_path))
            addMapping("Arial", 0, 0, "Arial")
            addMapping("Arial", 1, 0, "Arial")
            addMapping("Arial", 0, 1, "Arial")
            addMapping("Arial", 1, 1, "Arial")
            FONT_NAME = "Arial"
            logger.info("✔ Arial font loaded")
        else:
            logger.warning("⚠ Arial.ttf not found, fallback Helvetica")
            FONT_NAME = "Helvetica"
    except Exception as e:
        logger.error(f"Font load error: {e}")
        FONT_NAME = "Helvetica"

register_fonts()
addMapping("DejaVuSans", 0, 0, "DejaVuSans")
addMapping("DejaVuSans", 0, 1, "DejaVuSans")
addMapping("DejaVuSans", 1, 0, "DejaVuSans")
addMapping("DejaVuSans", 1, 1, "DejaVuSans")
# ==================== МОДЕЛИ ДАННЫХ ====================

class CityInfo(BaseModel):
    name: str
    population: int
    type: str
    region: str


class DynamicsResponse(BaseModel):
    city: str
    start_year: int
    end_year: int
    start_population: int
    end_population: int
    absolute_change: int
    relative_change: float
    cagr: float


class TopCitiesResponse(BaseModel):
    growing: List[DynamicsResponse]
    declining: List[DynamicsResponse]


class ForecastResponse(BaseModel):
    city: str
    last_year: int
    last_population: int
    horizon: int
    future_years: List[int]
    predictions: List[float]
    lower_bound: List[float]
    upper_bound: List[float]
    metrics: Optional[dict] = None


class MultiForecastRequest(BaseModel):
    cities: List[str]
    horizon: int = 10


class AIReportResponse(BaseModel):
    city: str
    region: str
    generated_at: str
    section_31_summary: str
    section_32_trends_and_factors: dict
    section_33_forecast: dict
    section_34_recommendations: List[dict]
    section_35_conclusion: str


class SaveForecastRequest(BaseModel):
    city: str
    horizon: int
    forecast_data: dict


class SaveForecastResponse(BaseModel):
    status: str
    id: int
    message: Optional[str] = None


class SavedForecastItem(BaseModel):
    id: int
    city: str
    horizon: int
    forecast_data: str
    created_at: str
    updated_at: str


class DeleteResponse(BaseModel):
    status: str
    message: str
    deleted_id: int


class ScenarioResponse(BaseModel):
    city: str
    last_population: int
    last_year: int
    scenarios: Dict[str, Dict[str, Any]]


# ==================== ИНИЦИАЛИЗАЦИЯ ====================

# Инициализируем базу данных
init_db()
logger.info("База данных инициализирована")

# Создаём экземпляры классов (они загрузят данные при старте)
monitor = PopulationMonitor()
forecaster = PopulationForecaster()
ai_analytics = AIAnalytics()
llm_analytics = LLMAnalytics()
# Создаём приложение FastAPI
app = FastAPI(
    title="Демографическая аналитика РФ",
    description="API для мониторинга, прогнозирования и AI-аналитики населения городов России",
    version="2.0.0",
    docs_url="/api/docs",
    redoc_url="/api/redoc"
)

# Монтируем статические файлы
static_dir = BASE_DIR / "static"
static_dir.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")
REPORTS_DIR = BASE_DIR / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

@app.on_event("startup")
def startup():
    REPORTS_DIR.mkdir(exist_ok=True)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ==================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ====================
def format_inline_markdown(text: str) -> str:
    """Преобразует **жирный**, *курсив* и _курсив_ в HTML-теги для ReportLab"""
    text = escape(text)
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = re.sub(r'_(.+?)_', r'<i>\1</i>', text)
    return text

def markdown_to_flowables(md_text: str, styles: dict) -> list:
    """Преобразует Markdown в список flowables (Paragraph, Spacer) для PDF"""
    flowables = []
    lines = md_text.splitlines()
    in_list = False

    for line in lines:
        line = line.strip()
        if not line:
            if in_list:
                in_list = False
            flowables.append(Spacer(1, 6))
            continue

        # Заголовки
        if line.startswith('# '):
            if in_list:
                in_list = False
            flowables.append(Paragraph(format_inline_markdown(line[2:]), styles['Heading1']))
            flowables.append(Spacer(1, 6))
        elif line.startswith('## '):
            if in_list:
                in_list = False
            flowables.append(Paragraph(format_inline_markdown(line[3:]), styles['Heading2']))
            flowables.append(Spacer(1, 6))
        elif line.startswith('### '):
            if in_list:
                in_list = False
            flowables.append(Paragraph(format_inline_markdown(line[4:]), styles['Heading3']))
            flowables.append(Spacer(1, 4))
        # Списки
        elif line.startswith('- ') or line.startswith('* '):
            text = format_inline_markdown(line[2:])
            flowables.append(Paragraph(f'• {text}', styles['Bullet']))
            in_list = True
        else:
            if in_list:
                in_list = False
                flowables.append(Spacer(1, 3))
            flowables.append(Paragraph(format_inline_markdown(line), styles['Normal']))
            flowables.append(Spacer(1, 6))

    return flowables
def convert_numpy_to_python(obj):
    """Рекурсивно преобразует numpy-типы в стандартные Python-типы для JSON-сериализации"""
    if isinstance(obj, (np.integer, np.int64, np.int32, np.int16, np.int8)):
        return int(obj)
    elif isinstance(obj, (np.floating, np.float64, np.float32, np.float16)):
        if np.isnan(obj) or np.isinf(obj):
            return None
        return float(round(obj, 6))
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, dict):
        return {key: convert_numpy_to_python(value) for key, value in obj.items()}
    elif isinstance(obj, (list, tuple)):
        return [convert_numpy_to_python(item) for item in obj]
    elif isinstance(obj, np.bool_):
        return bool(obj)
    elif isinstance(obj, bytes):
        return obj.decode('utf-8', errors='ignore')
    elif obj is None or isinstance(obj, (str, int, float, bool)):
        return obj
    else:
        try:
            return str(obj)
        except:
            return None


def sanitize_filename(filename: str) -> str:

    """Очистка имени файла для безопасного сохранения"""
    import re
    translit_map = {
        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'e',
        'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
        'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
        'ф': 'f', 'х': 'kh', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'shch',
        'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
        'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'E',
        'Ж': 'Zh', 'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M',
        'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U',
        'Ф': 'F', 'Х': 'Kh', 'Ц': 'Ts', 'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Shch',
        'Ъ': '', 'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya'
    }

    safe_name = ''.join(translit_map.get(c, c) for c in filename)
    safe_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', safe_name)
    safe_name = re.sub(r'_+', '_', safe_name)
    return safe_name.strip('_')



# ==================== БАЗОВЫЕ ЭНДПОИНТЫ ====================

@app.get("/ping")
async def ping():
    """Проверка работоспособности API"""
    return {
        "status": "ok",
        "message": "API демографической аналитики работает",
        "version": "2.0.0"
    }


@app.get("/")
async def root():
    """Главная страница"""
    return FileResponse(str(BASE_DIR / "index.html"))


# ==================== ЭНДПОИНТЫ МОНИТОРИНГА ====================

@app.get("/api/regions", response_model=List[str])
async def get_regions():
    """Получить список всех регионов"""
    all_cities = monitor.get_municipality_population()
    regions = sorted(set(city['region'] for city in all_cities if city.get('region')))
    return regions


@app.get("/api/cities", response_model=List[CityInfo])
async def get_cities(
    region: Optional[str] = None,
    city_type: Optional[str] = None,
    min_population: Optional[int] = None,
    limit: Optional[int] = None
):
    """Получить список городов с возможностью фильтрации"""
    cities = monitor.get_municipality_population()

    if region:
        cities = [c for c in cities if c.get('region') == region]
    if city_type:
        cities = [c for c in cities if c.get('type') == city_type]
    if min_population:
        cities = [c for c in cities if c.get('population', 0) >= min_population]

    if limit:
        cities = cities[:limit]

    return cities


@app.get("/api/cities/search")
async def search_cities(q: str, limit: int = 10):
    """Поиск городов по названию"""
    cities = monitor.get_municipality_population()
    q_lower = q.lower()
    matches = [c for c in cities if q_lower in c['name'].lower()]
    return matches[:limit]


@app.get("/api/dynamics/{city_name}", response_model=DynamicsResponse)
async def get_city_dynamics(city_name: str):
    """Получить динамику населения города за весь доступный период"""
    dynamics = monitor.get_population_dynamics(city_name)
    if dynamics is None:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден")
    return dynamics


@app.get("/api/top_changes", response_model=TopCitiesResponse)
async def get_top_changes(n: int = 10, min_population: int = 500000):
    """Получить топ городов по росту и снижению населения"""
    result = monitor.get_top_growing_declining(n=n, min_population=min_population)
    return result


@app.get("/api/regional_population")
async def get_regional_population():
    """Получить численность населения по регионам (для тепловой карты)"""
    regions = monitor.get_all_regions()
    return regions


@app.get("/api/demographics/{city_name}")
async def get_demographics(city_name: str):
    """Получить демографические показатели для города"""
    demo = monitor.get_demographic_indicators(city_name)
    if demo is None:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден")
    return demo


# ==================== ЭНДПОИНТЫ ПРОГНОЗИРОВАНИЯ ====================

@app.get("/api/forecast/{city_name}", response_model=ForecastResponse)
async def get_forecast(city_name: str, horizon: int = 10, include_metrics: bool = True):
    """Получить прогноз численности населения для города"""
    if horizon < 5 or horizon > 15:
        raise HTTPException(status_code=400, detail="Горизонт прогноза должен быть от 5 до 15 лет")

    forecast = forecaster.forecast_city(city_name, horizon=horizon)
    if forecast is None:
        raise HTTPException(
            status_code=404,
            detail=f"Город '{city_name}' не найден или недостаточно данных для прогноза"
        )

    response = {
        "city": forecast['city'],
        "last_year": int(forecast['last_year']),
        "last_population": int(forecast['last_population']),
        "horizon": horizon,
        "future_years": [int(y) for y in forecast['future_years']],
        "predictions": [float(p) for p in forecast['predictions']],
        "lower_bound": [float(lb) for lb in forecast['lower_bound']],
        "upper_bound": [float(ub) for ub in forecast['upper_bound']],
    }

    if include_metrics:
        metrics = forecaster.calculate_metrics(forecast)
        if metrics:
            response["metrics"] = metrics

    return response


@app.post("/api/forecast/compare")
async def compare_forecasts(request: MultiForecastRequest):
    """Сравнить прогнозы для нескольких городов"""
    results = {}
    for city in request.cities:
        forecast = forecaster.forecast_city(city, horizon=request.horizon)
        if forecast:
            results[city] = convert_numpy_to_python({
                "future_years": forecast['future_years'],
                "predictions": forecast['predictions'],
                "last_population": forecast['last_population']
            })
        else:
            results[city] = {"error": "Недостаточно данных"}
    return {"horizon": request.horizon, "results": results}


@app.get("/api/forecast/scenarios/{city_name}")
async def get_forecast_scenarios(city_name: str, horizon: int = 15):
    """Получить прогноз с тремя сценариями"""
    scenarios = forecaster.forecast_with_scenarios(city_name, horizon=horizon)
    if scenarios is None:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден")

    result = {
        "city": scenarios['city'],
        "last_population": int(scenarios['last_population']),
        "last_year": int(scenarios['last_year']),
        "scenarios": {}
    }

    for name, sc in scenarios['scenarios'].items():
        result["scenarios"][name] = {
            "rate": float(sc['rate']),
            "future_years": [int(y) for y in sc['future_years']],
            "predictions": [float(p) for p in sc['predictions']]
        }

    return result


# ==================== ЭНДПОИНТЫ AI-АНАЛИТИКИ ====================

@app.get("/api/ai/report/{city_name}")
async def get_ai_report(city_name: str):
    try:
        dynamics = monitor.get_population_dynamics(city_name)
        if dynamics is None:
            raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден")

        forecast = forecaster.forecast_city(city_name, horizon=10)
        if forecast is None:
            raise HTTPException(status_code=404, detail=f"Недостаточно данных для прогноза")

        # Получаем метрики отдельно, не добавляя их в forecast
        metrics = forecaster.calculate_metrics(forecast) if hasattr(forecaster, 'calculate_metrics') else None

        # Преобразуем numpy-типы в стандартные Python
        dynamics_clean = convert_numpy_to_python(dynamics)
        forecast_clean = convert_numpy_to_python(forecast)
        if metrics:
            metrics_clean = convert_numpy_to_python(metrics)
        else:
            metrics_clean = None

        # Передаём метрики отдельно, не изменяя forecast_clean
        report = llm_analytics.generate_report(city_name, dynamics_clean, forecast_clean, metrics_clean)
        return report
    except Exception as e:
        logger.error(f"LLM report error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Ошибка генерации AI-отчёта: {str(e)}")
@app.get("/api/ai/report/{city_name}/pdf")
async def get_ai_report_pdf(city_name: str):
    # 1. Получаем Markdown-отчёт
    md_content = ai_analytics.generate_markdown_report(city_name, forecast_horizon=10)
    if not md_content or "не найдены" in md_content:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден")

    # 2. Подготовка стилей
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.lib.styles import getSampleStyleSheet

    pdf_path = REPORTS_DIR / f"{sanitize_filename(city_name)}_report.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, leftMargin=72, rightMargin=72,
                            topMargin=72, bottomMargin=72)

    base_styles = getSampleStyleSheet()
    styles = {}

    # Обычный текст
    styles['Normal'] = ParagraphStyle(
        'CustomNormal', parent=base_styles['Normal'],
        fontName=FONT_NAME, fontSize=11, leading=14, spaceAfter=6
    )
    # Заголовок 1 уровня
    styles['Heading1'] = ParagraphStyle(
        'CustomHeading1', parent=base_styles['Heading1'],
        fontName=FONT_NAME, fontSize=18, textColor=colors.HexColor("#1f4e79"),
        spaceAfter=12, spaceBefore=6, bold=True
    )
    # Заголовок 2 уровня
    styles['Heading2'] = ParagraphStyle(
        'CustomHeading2', parent=base_styles['Heading2'],
        fontName=FONT_NAME, fontSize=14, textColor=colors.HexColor("#2e6b9e"),
        spaceAfter=8, spaceBefore=4, bold=True
    )
    # Заголовок 3 уровня
    styles['Heading3'] = ParagraphStyle(
        'CustomHeading3', parent=base_styles['Heading3'],
        fontName=FONT_NAME, fontSize=12, textColor=colors.HexColor("#4a86b8"),
        spaceAfter=6, spaceBefore=2, bold=True
    )
    # Список (буллеты)
    styles['Bullet'] = ParagraphStyle(
        'CustomBullet', parent=styles['Normal'],
        leftIndent=20, firstLineIndent=0, bulletIndent=10, spaceAfter=3
    )

    # 3. Преобразуем Markdown в flowables
    story = markdown_to_flowables(md_content, styles)

    # 4. Строим PDF
    doc.build(story)

    # 5. Отдаём файл
    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        filename=f"{sanitize_filename(city_name)}_report.pdf"
    )
@app.get("/api/ai/report/{city_name}/docx")
async def get_ai_report_docx(city_name: str):
    """Скачать аналитический отчёт в формате Word (DOCX)"""
    md_content = ai_analytics.generate_markdown_report(city_name, forecast_horizon=10)
    if not md_content or "не найдены" in md_content:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден в данных")

    doc = Document()
    doc.add_heading(f"Аналитическая справка: {city_name}", 0)

    for line in md_content.split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    safe_name = sanitize_filename(city_name)
    file_path = REPORTS_DIR / f"{safe_name}_report.docx"
    doc.save(str(file_path))

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{safe_name}_report.docx"
    )
@app.get("/api/ai/report/{city_name}/markdown")
async def get_ai_report_markdown(city_name: str):
    """Скачать аналитический отчёт в формате Markdown"""
    md_content = ai_analytics.generate_markdown_report(city_name, forecast_horizon=10)

    if md_content is None or "не найдены" in md_content:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден в данных")

    safe_name = sanitize_filename(city_name)

    return PlainTextResponse(
        content=md_content,
        media_type="text/markdown; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename={safe_name}_analytics_report.md"
        }
    )


@app.get("/api/ai/summary/{city_name}")
async def get_ai_summary(city_name: str):
    """Получить только краткое резюме по городу"""
    summary = ai_analytics.generate_summary(city_name)
    if summary is None or "не найдены" in summary:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден")
    return {"city": city_name, "summary": summary}


@app.get("/api/ai/recommendations/{city_name}")
async def get_ai_recommendations(city_name: str):
    """Получить только рекомендации по городу"""
    recommendations = ai_analytics.generate_recommendations(city_name)
    if not recommendations:
        raise HTTPException(status_code=404, detail=f"Город '{city_name}' не найден или нет рекомендаций")

    recommendations = convert_numpy_to_python(recommendations)

    return {
        "city": city_name,
        "total_recommendations": len(recommendations),
        "recommendations": recommendations
    }


# ==================== ЭНДПОИНТЫ БАЗЫ ДАННЫХ ====================

@app.post("/api/forecasts/save", response_model=SaveForecastResponse)
async def save_forecast_endpoint(request: SaveForecastRequest):
    """Сохранить прогноз в базу данных"""
    try:
        if not request.city:
            raise HTTPException(status_code=400, detail="Название города обязательно")
        if request.horizon < 1 or request.horizon > 50:
            raise HTTPException(status_code=400, detail="Горизонт прогноза должен быть от 1 до 50 лет")

        record_id = save_forecast(
            city=request.city,
            horizon=request.horizon,
            forecast_data=request.forecast_data
        )

        logger.info(f"Прогноз для {request.city} сохранен (id={record_id})")

        return {
            "status": "ok",
            "id": record_id,
            "message": f"Прогноз для города {request.city} успешно сохранен"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Ошибка сохранения прогноза: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка сохранения: {str(e)}")


@app.get("/api/forecasts/saved")
async def get_saved_forecasts(city: Optional[str] = None, limit: Optional[int] = None):
    """Получить список сохраненных прогнозов"""
    try:
        if city:
            forecasts = get_forecasts_by_city(city)
        else:
            forecasts = get_all_forecasts()

        if limit:
            forecasts = forecasts[:limit]

        return forecasts
    except Exception as e:
        logger.error(f"Ошибка получения прогнозов: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка получения данных: {str(e)}")


@app.get("/api/forecasts/saved/{forecast_id}")
async def get_saved_forecast(forecast_id: int):
    """Получить сохраненный прогноз по ID"""
    forecast = get_forecast_by_id(forecast_id)
    if forecast is None:
        raise HTTPException(status_code=404, detail=f"Прогноз с ID {forecast_id} не найден")
    return forecast


@app.delete("/api/forecasts/saved/{forecast_id}", response_model=DeleteResponse)
async def delete_saved_forecast(forecast_id: int):
    """Удалить сохраненный прогноз"""
    success = delete_forecast(forecast_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"Прогноз с ID {forecast_id} не найден")

    logger.info(f"Прогноз {forecast_id} удален")

    return {
        "status": "ok",
        "message": f"Прогноз {forecast_id} успешно удален",
        "deleted_id": forecast_id
    }


# ==================== ДОПОЛНИТЕЛЬНЫЕ ЭНДПОИНТЫ ====================

@app.get("/api/available_years")
async def get_available_years():
    """Получить диапазон лет, за которые есть исторические данные"""
    years = sorted(monitor.historical['year'].unique())
    return {
        "min_year": int(min(years)),
        "max_year": int(max(years)),
        "all_years": [int(y) for y in years],
        "total_years": len(years)
    }


@app.get("/api/statistics")
async def get_statistics():
    """Получить общую статистику по данным"""
    try:
        last_year = monitor.historical['year'].max()
        data = monitor.historical[monitor.historical['year'] == last_year]

        total_population = data['population'].sum()
        cities_count = len(data)
        avg_population = data['population'].mean()

        million_cities = data[data['population'] >= 1_000_000]
        large_cities = data[(data['population'] >= 500_000) & (data['population'] < 1_000_000)]

        saved_forecasts = get_all_forecasts()

        return {
            "year": int(last_year),
            "total_population": int(total_population),
            "total_cities": cities_count,
            "average_population": float(avg_population),
            "million_cities": len(million_cities),
            "large_cities": len(large_cities),
            "saved_forecasts_count": len(saved_forecasts)
        }
    except Exception as e:
        logger.error(f"Ошибка получения статистики: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка получения статистики: {str(e)}")
